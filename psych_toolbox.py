"""
团辅工具箱
整合：团辅策划案生成器 + 团辅报销助手
部署到 Streamlit Community Cloud
"""

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import os
import io
import sys
from PIL import Image
import pdfplumber

st.set_page_config(page_title="团辅工具箱", page_icon="\U0001f3e5", layout="wide")

# ===== 侧边栏导航 =====
st.sidebar.title("\U0001f3e5 团辅工具箱")
tool = st.sidebar.radio("选择工具", ["\U0001f4dd 团辅策划案生成器", "\U0001f4cb 团辅报销助手"], index=0)

st.markdown("""
<style>
@media (max-width: 768px) {
    .stColumns {
        flex-direction: column !important;
    }
    [data-testid="stSidebar"] {
        width: 100% !important;
    }
    .stTextArea textarea {
        font-size: 16px !important;
    }
    .stTextInput input {
        font-size: 16px !important;
    }
    .stSelectbox select {
        font-size: 16px !important;
    }
}
</style>
""", unsafe_allow_html=True)


# ===== 通用函数 =====

def add_docx_to_doc(source_docx_bytes, target_doc, max_width=Inches(6)):
    """将上传的docx内容（文字+图片）复制到目标文档"""
    src = Document(io.BytesIO(source_docx_bytes))
    
    # 复制段落
    for para in src.paragraphs:
        text = para.text.strip()
        if not text:
            # 空段落，添加空行
            target_doc.add_paragraph()
            continue
        
        p = target_doc.add_paragraph()
        # 尝试复制对齐方式
        if para.alignment is not None:
            p.alignment = para.alignment
        
        for run in para.runs:
            r = p.add_run(run.text)
            r.bold = run.bold
            r.italic = run.italic
            if run.font.size:
                r.font.size = run.font.size
            if run.font.name:
                r.font.name = run.font.name
                rPr = r._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = r._element.makeelement(qn('w:rFonts'), {})
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), run.font.name)
    
    # 复制图片
    for rel in src.part.rels.values():
        if "image" in rel.reltype:
            image = rel.target_part
            img_bytes = image.blob
            img = Image.open(io.BytesIO(img_bytes))
            
            # 获取原始尺寸（EMU单位）
            width_emu = None
            for rel2 in src.part.rels.values():
                if rel2.target_part == image:
                    for blip in src.element.body.iter(qn('wp:inline')):
                        ext = blip.find(qn('wp:extent'))
                        if ext is not None:
                            width_emu = int(ext.get('cx', 0))
                            break
                    break
            
            if width_emu and width_emu > 0:
                width_inches = width_emu / 914400
                if width_inches > 6:
                    width_inches = 6
                target_doc.add_picture(io.BytesIO(img_bytes), width=Inches(width_inches))
            else:
                target_doc.add_picture(io.BytesIO(img_bytes), width=max_width)


def add_pdf_to_doc(pdf_bytes, target_doc, max_width=Inches(6)):
    """将PDF每一页渲染为图片添加到目标文档"""
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            img = page.to_image(resolution=200)
            pil_image = img.original
            
            width, height = pil_image.size
            max_width_px = int(max_width / Inches(1) * 96)
            if width > max_width_px:
                ratio = max_width_px / width
                new_size = (int(width * ratio), int(height * ratio))
                pil_image = pil_image.resize(new_size, Image.Resampling.LANCZOS)
            
            temp_path = f"/tmp/pdf_page_{page_num}.png"
            pil_image.save(temp_path)
            target_doc.add_picture(temp_path, width=max_width)
            if os.path.exists(temp_path):
                os.remove(temp_path)


# ============================================================
#  团辅策划案生成器
# ============================================================
if tool == "\U0001f4dd 团辅策划案生成器":

    st.title("\U0001f4dd 团辅策划案生成器")
    st.markdown("填写内容，生成标准格式的策划案文档")
    st.markdown("---")

    col_form, col_ref = st.columns([3, 2])

    with col_form:
        with st.form("planning_form"):
            st.subheader("基本信息")

            activity_date = st.text_input("活动时间", placeholder="请参考样文")

            loc_sel = st.selectbox("活动地点", ["暨南大学番禺校区知识产权大楼802室", "自定义地点"])
            if loc_sel == "自定义地点":
                activity_location = st.text_input("自定义活动地点")
            else:
                activity_location = loc_sel

            aud_sel = st.selectbox("参与对象", ["暨南大学番禺校区在校学生", "自定义对象"])
            if aud_sel == "自定义对象":
                target_audience = st.text_input("自定义参与对象")
            else:
                target_audience = aud_sel

            participant_count = st.text_input("人数", placeholder="请参考样文")

            st.subheader("活动内容")

            st.write("**活动标题**")
            activity_title = st.text_area("活动标题", height=80, placeholder="请参考样文")

            st.write("**活动背景**")
            activity_background = st.text_area("活动背景内容", height=120, placeholder="请参考样文")

            st.write("**活动目的**")
            activity_purpose = st.text_area("活动目的内容", height=120, placeholder="请参考样文")

            st.write("**活动主题**")
            activity_theme = st.text_area("活动主题", height=100, placeholder="请参考样文")

            st.write("**活动准备**")
            activity_prep = st.text_area("活动准备内容", height=120, placeholder="请参考样文")
            st.caption("格式：每行一项，用空格或 | 分隔三列：准备项目 具体内容 备注")

            st.write("**活动流程**")
            activity_flow = st.text_area("活动流程内容", height=150, placeholder="请参考样文")
            st.caption("格式：每行一个环节，用空格或 | 分隔三列：时间 事项 备注")

            st.write("**活动物资及预算**")
            activity_materials = st.text_area("物资及预算内容", height=120, placeholder="请参考样文")
            st.caption("格式：每行一个物品，用空格或 | 分隔五列：项目 单位 单价 数量 备注（序号和项目总计自动生成）")

            st.write("**注意事项及应急措施**")
            activity_notes = st.text_area("注意事项内容", height=120, placeholder="请参考样文")
            st.caption("格式：每行一条，用空格或 | 分隔两列：类别 内容")

            submitted = st.form_submit_button("生成策划案", type="primary", use_container_width=True)

    with col_ref:
        st.subheader("\U0001f4d6 参考示例")
        ref_images = []
        for i in range(1, 7):
            img_path = f'reference_page_{i}.png'
            if os.path.exists(img_path):
                ref_images.append(img_path)
        if ref_images:
            for i, img_path in enumerate(ref_images, 1):
                with st.expander(f"参考文档第 {i} 页"):
                    st.image(img_path, use_container_width=True)
        else:
            st.warning("参考图片未找到")

    # ===== 解析函数 =====
    def parse_prep_to_table(prep_text):
        lines = prep_text.strip().split('\n')
        table_data = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if '|' in line:
                parts = [p.strip() for p in line.split('|')]
            else:
                parts = [p.strip() for p in line.split(None, 2)]
            if len(parts) >= 3:
                table_data.append(parts[:3])
            elif len(parts) == 2:
                table_data.append([parts[0], parts[1], ''])
            elif len(parts) == 1:
                table_data.append(['', parts[0], ''])
        return table_data

    def parse_flow_to_table(flow_text):
        lines = flow_text.strip().split('\n')
        table_data = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if '|' in line:
                parts = [p.strip() for p in line.split('|')]
            else:
                parts = [p.strip() for p in line.split(None, 2)]
            if len(parts) >= 3:
                table_data.append(parts[:3])
            elif len(parts) == 2:
                table_data.append([parts[0], parts[1], ''])
            elif len(parts) == 1:
                table_data.append(['', parts[0], ''])
        return table_data

    def parse_materials_to_table(materials_text):
        lines = materials_text.strip().split('\n')
        table_data = []
        for idx, line in enumerate(lines, 1):
            line = line.strip()
            if not line:
                continue
            if '|' in line:
                parts = [p.strip() for p in line.split('|')]
            else:
                parts = [p.strip() for p in line.split(None, 4)]
            item = parts[0] if len(parts) > 0 else ''
            unit = parts[1] if len(parts) > 1 else ''
            price = parts[2] if len(parts) > 2 else '0'
            qty = parts[3] if len(parts) > 3 else '1'
            note = parts[4] if len(parts) > 4 else ''
            try:
                total = str(float(qty) * float(price))
            except:
                total = '0'
            table_data.append([str(idx), item, unit, price, qty, total, note])
        return table_data

    def parse_notes_to_list(notes_text):
        lines = notes_text.strip().split('\n')
        items = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if '|' in line:
                parts = [p.strip() for p in line.split('|')]
            else:
                parts = [p.strip() for p in line.split(None, 1)]
            if len(parts) >= 2:
                items.append((parts[0], parts[1]))
            elif len(parts) == 1:
                items.append(('', parts[0]))
        return items

    # ===== 生成文档 =====
    if submitted:
        with st.spinner("正在生成策划案..."):
            try:
                title_text = activity_title.strip() if activity_title.strip() else "\u201c\u91cd\u5851\u81ea\u6211\uff1a\u753b\u7b14\u4e0e\u7c98\u571f\u7684\u6210\u957f\u63a2\u7d22\u201d\u5fc3\u7406\u56e2\u8f85\u6d3b\u52a8\u7b56\u5212\u6848"
                theme_text = activity_theme.strip() if activity_theme.strip() else "待定"
                background_text = activity_background.strip() if activity_background.strip() else "待定"
                purpose_text = activity_purpose.strip() if activity_purpose.strip() else "待定"
                prep_text = activity_prep.strip() if activity_prep.strip() else ""
                flow_text = activity_flow.strip() if activity_flow.strip() else ""
                materials_text = activity_materials.strip() if activity_materials.strip() else ""
                notes_text = activity_notes.strip() if activity_notes.strip() else ""

                prep_table = parse_prep_to_table(prep_text) if prep_text else []
                flow_table = parse_flow_to_table(flow_text) if flow_text else []
                materials_table = parse_materials_to_table(materials_text) if materials_text else []
                notes_list = parse_notes_to_list(notes_text) if notes_text else []

                doc = Document()

                style = doc.styles['Normal']
                style.font.name = '仿宋_GB2312'
                style.font.size = Pt(14)
                style.font.color.rgb = RGBColor(0, 0, 0)
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(3)
                style.paragraph_format.line_spacing = 1.5
                style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

                def set_font(run, size=14, bold=False):
                    run.font.name = '仿宋_GB2312'
                    run.font.size = Pt(size)
                    run.font.bold = bold
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = run._element.makeelement(qn('w:rFonts'), {})
                        rPr.insert(0, rFonts)
                    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

                def add_body(text):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.first_line_indent = Pt(28)
                    r = p.add_run(text)
                    set_font(r, 14)
                    return p

                def add_section_title(text):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(6)
                    r = p.add_run(text)
                    set_font(r, 16, bold=True)
                    return p

                def add_sub_title(text):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(3)
                    r = p.add_run(text)
                    set_font(r, 14, bold=True)
                    return p

                def add_sub_sub_title(text):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    r = p.add_run(text)
                    set_font(r, 14, bold=True)
                    return p

                def style_cell(cell, text, size=12, bold=False, center=False):
                    cell.text = ''
                    p = cell.paragraphs[0]
                    r = p.add_run(text)
                    set_font(r, size, bold)
                    if center:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(12)
                r = p.add_run(title_text)
                set_font(r, 22, bold=True)

                add_section_title('一、活动背景')
                add_body(background_text)

                add_section_title('二、活动目的')
                add_body(purpose_text)

                add_section_title('三、活动简介')
                add_sub_title('（一）活动主题')
                add_body(theme_text)
                add_sub_title('（二）活动时间')
                add_body(activity_date if activity_date else "待定")
                add_sub_title('（三）活动地点')
                add_body(activity_location if activity_location else "待定")
                add_sub_title('（四）参与对象')
                audience_text = f'{target_audience if target_audience else "待定"}（人数：{participant_count if participant_count else "待定"}）'
                add_body(audience_text)

                add_section_title('四、活动准备')
                if prep_table:
                    tbl = doc.add_table(rows=len(prep_table)+1, cols=3)
                    tbl.style = 'Table Grid'
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    headers = ['准备项目', '具体内容', '备注']
                    for i, h in enumerate(headers):
                        style_cell(tbl.rows[0].cells[i], h, 12, True, True)
                    for row_idx, row_data in enumerate(prep_table, 1):
                        for col_idx in range(3):
                            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
                            style_cell(tbl.rows[row_idx].cells[col_idx], cell_text, 12)

                add_section_title('五、活动执行')
                add_sub_sub_title('1、活动流程')
                if flow_table:
                    tbl = doc.add_table(rows=len(flow_table)+1, cols=3)
                    tbl.style = 'Table Grid'
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    headers = ['时间', '事项', '备注']
                    for i, h in enumerate(headers):
                        style_cell(tbl.rows[0].cells[i], h, 12, True, True)
                    for row_idx, row_data in enumerate(flow_table, 1):
                        for col_idx in range(3):
                            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
                            style_cell(tbl.rows[row_idx].cells[col_idx], cell_text, 12)

                add_section_title('六、活动物资及预算')
                if materials_table:
                    tbl = doc.add_table(rows=len(materials_table)+1, cols=7)
                    tbl.style = 'Table Grid'
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    headers = ['序号', '项目', '单位', '单价', '数量', '项目总计', '备注']
                    for i, h in enumerate(headers):
                        style_cell(tbl.rows[0].cells[i], h, 12, True, True)
                    for row_idx, row_data in enumerate(materials_table, 1):
                        for col_idx in range(7):
                            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
                            style_cell(tbl.rows[row_idx].cells[col_idx], cell_text, 12)

                add_section_title('七、注意事项及应急措施')
                circled = ['\u2460', '\u2461', '\u2462', '\u2463', '\u2464', '\u2465', '\u2466', '\u2467', '\u2468', '\u2469']
                for idx, (cat, content) in enumerate(notes_list):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.first_line_indent = Pt(28)
                    num = circled[idx] if idx < len(circled) else f'({idx+1})'
                    if cat:
                        r = p.add_run(f'{num}  {cat}：')
                        set_font(r, 14, bold=True)
                        r = p.add_run(f'  {content}')
                        set_font(r, 14)
                    else:
                        r = p.add_run(f'{num}  {content}')
                        set_font(r, 14)

                output_filename = f"策划案_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                doc.save(output_filename)

                with open(output_filename, "rb") as file:
                    st.download_button(
                        label="\U0001f4e5 下载策划案",
                        data=file,
                        file_name=output_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

                st.success("\u2705 策划案生成成功！")
                os.remove(output_filename)

            except Exception as e:
                st.error(f"生成策划案时出错：{str(e)}")
                import traceback
                st.code(traceback.format_exc())

    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: gray; font-size: 12px;'>
        团辅策划案生成器 | 纯工具模式，不存储填写内容
        </div>
        """,
        unsafe_allow_html=True
    )

# ============================================================
#  团辅报销助手
# ============================================================
elif tool == "\U0001f4cb 团辅报销助手":

    st.title("\U0001f4cb 团辅报销助手")
    st.markdown("---")

    col1, col2 = st.columns(2)

    with col1:
        planning_file = st.file_uploader(
            "1. 上传策划案（Word 或 PDF）",
            type=['docx', 'pdf'],
            help="支持 .docx 和 .pdf 格式"
        )
        purchase_images = st.file_uploader(
            "2. 上传购买截图（可多选）",
            type=['png', 'jpg', 'jpeg'],
            accept_multiple_files=True,
            help="支持 PNG、JPG 格式"
        )
        activity_doc = st.file_uploader(
            "5. 上传活动文字说明（Word 或 PDF）",
            type=['docx', 'pdf'],
            help="支持 .docx 和 .pdf 格式"
        )

    with col2:
        invoice_files = st.file_uploader(
            "3. 上传发票（图片或 PDF，可多选）",
            type=['png', 'jpg', 'jpeg', 'pdf'],
            accept_multiple_files=True,
            help="支持 PNG、JPG、PDF 格式"
        )
        activity_images = st.file_uploader(
            "4. 上传活动照片（最多 9 张）",
            type=['png', 'jpg', 'jpeg'],
            accept_multiple_files=True,
            help="最多 9 张，支持 PNG、JPG 格式"
        )

    st.markdown("---")
    if st.button("\U0001f680 生成报销文档", type="primary", use_container_width=True):

        if not planning_file:
            st.error("请上传策划案文件")
        else:
            try:
                with st.spinner("正在生成文档..."):
                    doc = Document()

                    style = doc.styles['Normal']
                    font = style.font
                    font.name = '宋体'
                    font.size = Pt(12)

                    def resize_image_smart(img, max_width_inches=6, max_height_inches=8):
                        dpi = 96
                        max_width_px = int(max_width_inches * dpi)
                        max_height_px = int(max_height_inches * dpi)
                        width, height = img.size
                        width_ratio = max_width_px / width if width > max_width_px else 1
                        height_ratio = max_height_px / height if height > max_height_px else 1
                        ratio = min(width_ratio, height_ratio)
                        if height > width * 1.5:
                            ratio = min(ratio, 0.6)
                        if ratio < 1:
                            new_size = (int(width * ratio), int(height * ratio))
                            img = img.resize(new_size, Image.Resampling.LANCZOS)
                        return img

                    # ===== 第一部分：策划案 =====
                    heading = doc.add_heading('活动策划案', level=1)
                    heading.paragraph_format.space_before = Pt(0)
                    heading.paragraph_format.space_after = Pt(6)
                    heading.paragraph_format.keep_with_next = True

                    if planning_file.type == "application/pdf":
                        add_pdf_to_doc(planning_file.read(), doc)
                    elif planning_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        add_docx_to_doc(planning_file.read(), doc)

                    doc.add_page_break()

                    # ===== 第二部分：购买截图 =====
                    if purchase_images:
                        heading = doc.add_heading('购买截图', level=1)
                        heading.paragraph_format.space_before = Pt(0)
                        heading.paragraph_format.space_after = Pt(6)
                        heading.paragraph_format.keep_with_next = True
                        for idx, img_file in enumerate(purchase_images, 1):
                            img = Image.open(img_file)
                            img = resize_image_smart(img)
                            actual_width = Inches(img.width / 96)
                            temp_path = f"/tmp/purchase_{idx}.png"
                            img.save(temp_path)
                            doc.add_picture(temp_path, width=actual_width)
                            doc.add_paragraph(f"购买截图 {idx}", style='Caption')
                            if os.path.exists(temp_path):
                                os.remove(temp_path)
                        doc.add_page_break()

                    # ===== 第三部分：发票 =====
                    if invoice_files:
                        heading = doc.add_heading('发票', level=1)
                        heading.paragraph_format.space_before = Pt(0)
                        heading.paragraph_format.space_after = Pt(6)
                        heading.paragraph_format.keep_with_next = True
                        invoice_idx = 1
                        for file in invoice_files:
                            if file.type == "application/pdf":
                                with pdfplumber.open(io.BytesIO(file.read())) as pdf:
                                    for page_num, page in enumerate(pdf.pages):
                                        img = page.to_image(resolution=200)
                                        pil_image = img.original
                                        max_width = Inches(6)
                                        width, height = pil_image.size
                                        max_width_px = int(max_width / Inches(1) * 96)
                                        if width > max_width_px:
                                            ratio = max_width_px / width
                                            new_size = (int(width * ratio), int(height * ratio))
                                            pil_image = pil_image.resize(new_size, Image.Resampling.LANCZOS)
                                        temp_path = f"/tmp/invoice_pdf_{invoice_idx}.png"
                                        pil_image.save(temp_path)
                                        doc.add_picture(temp_path, width=max_width)
                                        doc.add_paragraph(f"发票 {invoice_idx}（第 {page_num + 1} 页）", style='Caption')
                                        if os.path.exists(temp_path):
                                            os.remove(temp_path)
                                        invoice_idx += 1
                            else:
                                img = Image.open(file)
                                img = resize_image_smart(img)
                                actual_width = Inches(img.width / 96)
                                temp_path = f"/tmp/invoice_{invoice_idx}.png"
                                img.save(temp_path)
                                doc.add_picture(temp_path, width=actual_width)
                                doc.add_paragraph(f"发票 {invoice_idx}", style='Caption')
                                if os.path.exists(temp_path):
                                    os.remove(temp_path)
                                invoice_idx += 1
                        doc.add_page_break()

                    # ===== 第四部分：活动照片 =====
                    if activity_images:
                        activity_images = activity_images[:9]
                        heading = doc.add_heading('活动照片', level=1)
                        heading.paragraph_format.space_before = Pt(0)
                        heading.paragraph_format.space_after = Pt(6)
                        heading.paragraph_format.keep_with_next = True
                        for idx, img_file in enumerate(activity_images, 1):
                            img = Image.open(img_file)
                            img = resize_image_smart(img)
                            actual_width = Inches(img.width / 96)
                            temp_path = f"/tmp/activity_{idx}.png"
                            img.save(temp_path)
                            doc.add_picture(temp_path, width=actual_width)
                            doc.add_paragraph(f"活动照片 {idx}", style='Caption')
                            if os.path.exists(temp_path):
                                os.remove(temp_path)

                    # ===== 第五部分：活动文字说明 =====
                    if activity_doc:
                        doc.add_page_break()
                        heading = doc.add_heading('活动文字说明', level=1)
                        heading.paragraph_format.space_before = Pt(0)
                        heading.paragraph_format.space_after = Pt(6)
                        heading.paragraph_format.keep_with_next = True

                        if activity_doc.type == "application/pdf":
                            add_pdf_to_doc(activity_doc.read(), doc)
                        elif activity_doc.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            add_docx_to_doc(activity_doc.read(), doc)

                    # 保存
                    output_filename = f"报销材料_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    doc.save(output_filename)

                    with open(output_filename, "rb") as file:
                        st.download_button(
                            label="\U0001f4e5 下载报销文档",
                            data=file,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                    st.success("\u2705 文档生成成功！")
                    os.remove(output_filename)

            except Exception as e:
                st.error(f"生成文档时出错：{str(e)}")
                import traceback
                st.code(traceback.format_exc())

    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: gray; font-size: 12px;'>
        团辅报销助手 | 纯工具模式，不存储上传文件
        </div>
        """,
        unsafe_allow_html=True
    )
